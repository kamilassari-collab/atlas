"""
scripts/render_cv.py — produce the tailored CV for a job.

Two paths, per the sacred CV rules (design doc, CEO review D2):
  DOCX user  → swap bullets IN the user's own DOCX (design untouched).
               Usage: uv run python -m scripts.render_cv <job_id> '<bullets json>'
               bullets json: {"Role heading text": ["bullet 1", "bullet 2", ...], ...}
  PDF/text   → we never rebuild; write a suggestions markdown the user applies
               in their own file.
               Usage: uv run python -m scripts.render_cv <job_id> --suggestions '<json>'
               suggestions json: [{"original": "...", "suggestion": "..."}, ...]

Writes into data/outputs/, updates the DB (tailored_cv_path + status=cv_tailored),
prints the saved path.
"""
from __future__ import annotations
import json
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import yaml
from pipeline import db

ROOT = Path(__file__).parent.parent
OUT = ROOT / "data" / "outputs"
CONFIG = ROOT / "config.yaml"


def _cfg() -> dict:
    if CONFIG.exists():
        return yaml.safe_load(CONFIG.read_text()) or {}
    return {}


def _safe(s: str) -> str:
    return re.sub(r"[^\w]+", "_", s)[:24].strip("_") or "x"


# ── DOCX → PDF conversion (preserve the user's design, deliver a PDF) ─────────

def _find_soffice() -> str | None:
    """Locate a LibreOffice/soffice binary (free, cross-platform, headless)."""
    import shutil
    for name in ("soffice", "libreoffice"):
        p = shutil.which(name)
        if p:
            return p
    for p in (
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/soffice", "/usr/bin/libreoffice",
        "/opt/homebrew/bin/soffice",
    ):
        if Path(p).exists():
            return p
    return None


def docx_to_pdf(docx_path: Path) -> Path | None:
    """Convert a DOCX to PDF in place (same folder), preserving its design.

    Tries the free LibreOffice headless path first, then docx2pdf (which drives
    MS Word, if installed). Returns the PDF path, or None if no converter is
    available — callers then keep the DOCX and tell the user how to get PDFs.
    """
    import subprocess
    pdf_path = docx_path.with_suffix(".pdf")

    soffice = _find_soffice()
    if soffice:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(docx_path.parent), str(docx_path)],
                check=True, capture_output=True, timeout=120,
            )
            if pdf_path.exists():
                return pdf_path
        except (subprocess.SubprocessError, OSError):
            pass

    try:  # docx2pdf uses MS Word (macOS/Windows) — works on machines with Word.
        from docx2pdf import convert as _convert
        _convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass

    return None


# ── DOCX path: in-place bullet swap (ported from job_machine) ────────────────

def swap_docx(job_id: int, tailored: dict) -> Path:
    import docx as _docx

    job = db.get_job(job_id)
    if job is None:
        raise SystemExit(f"No job with id {job_id}")
    cv_cfg = _cfg().get("cv", {}) or {}
    src = cv_cfg.get("docx_path")
    if not src or not Path(src).exists():
        raise SystemExit("cv.docx_path missing in config.yaml — is the CV a DOCX?")

    doc = _docx.Document(src)
    tailored_norm = {k.lower().strip(): v for k, v in tailored.items()}
    used_keys: set[str] = set()

    def match_role(text: str):
        """Match the heading to the single best key.

        The agent passes keys as role headings copied verbatim from the CV, so
        prefer exact/substring matches. Word-overlap is only a last resort and
        requires the heading to share ALL of the key's significant words — not
        just one — so 'Chief of Staff Intern' never bleeds into every other
        '... Intern' heading. Each key is consumed once, so two similar roles
        can't both grab the same bullets.
        """
        t = text.lower().strip()
        # 1. exact, then containment either direction (most specific wins).
        candidates = [k for k in tailored_norm if k not in used_keys]
        for k in sorted(candidates, key=len, reverse=True):
            if k == t or k in t or t in k:
                used_keys.add(k)
                return tailored_norm[k]
        # 2. last resort: heading contains ALL significant (>4 char) key words.
        for k in sorted(candidates, key=len, reverse=True):
            sig = [w for w in k.split() if len(w) > 4]
            if sig and all(w in t for w in sig):
                used_keys.add(k)
                return tailored_norm[k]
        return None

    current_bullets: list = []
    idx = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            current_bullets, idx = [], 0
            continue
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        style = para.style.name.lower()
        is_bullet = "list" in style or text.startswith(("•", "-"))

        if is_bold and not is_bullet:
            matched = match_role(text)
            if matched:
                current_bullets, idx = matched[:], 0
        elif is_bullet and current_bullets and idx < len(current_bullets):
            new_text = current_bullets[idx]
            idx += 1
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

    OUT.mkdir(parents=True, exist_ok=True)
    docx_path = OUT / f"cv_{job_id}_{_safe(job.company)}.docx"
    doc.save(str(docx_path))

    # The bullet swap is instant — link the DOCX NOW so the card lights up
    # immediately, then upgrade it to a PDF in a detached background process.
    # No blocking on Word/LibreOffice; the user is never left waiting.
    rel_docx = str(docx_path.relative_to(ROOT))
    db.update_job(job_id, tailored_cv_path=rel_docx, status="cv_tailored")

    import subprocess
    subprocess.Popen(
        [sys.executable, "-m", "scripts.render_cv", "__convert__", str(job_id), rel_docx],
        cwd=str(ROOT), stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        start_new_session=True,
    )
    return docx_path


# ── PDF/text path: suggestions list, never a rebuild ─────────────────────────

def write_suggestions(job_id: int, suggestions: list) -> Path:
    """Render a suggestions PDF for PDF/text CVs — never rebuilds the CV itself.

    We don't have the user's PDF design, so we don't recreate their CV (sacred
    rule). Instead we deliver a clean PDF list of "replace X with Y" edits they
    apply in their own file. Output is PDF so every job gets a PDF deliverable.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT

    job = db.get_job(job_id)
    if job is None:
        raise SystemExit(f"No job with id {job_id}")
    OUT.mkdir(parents=True, exist_ok=True)
    out_path = OUT / f"cv_suggestions_{job_id}_{_safe(job.company)}.pdf"

    esc = lambda s: (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    h1 = ParagraphStyle("h1", fontName="Helvetica-Bold", fontSize=15, leading=19, spaceAfter=4)
    intro = ParagraphStyle("intro", fontName="Helvetica", fontSize=10, leading=15,
                           textColor="#555555", spaceAfter=14)
    num = ParagraphStyle("num", fontName="Helvetica-Bold", fontSize=10.5, leading=15,
                         textColor="#5B4CF5", spaceBefore=8, spaceAfter=2)
    old = ParagraphStyle("old", fontName="Helvetica", fontSize=10, leading=15, alignment=TA_LEFT)
    new = ParagraphStyle("new", fontName="Helvetica-Bold", fontSize=10, leading=15,
                         alignment=TA_LEFT, spaceAfter=6)

    story = [
        Paragraph(f"Suggestions CV — {esc(job.title)} @ {esc(job.company)}", h1),
        Paragraph("Ton CV reste TON fichier : applique ces remplacements toi-même. "
                  "Chaque suggestion ré-angle ton expérience réelle — rien d'inventé.", intro),
    ]
    for i, s in enumerate(suggestions, 1):
        story += [
            Paragraph(f"{i}.", num),
            Paragraph(f"<font color='#999999'>Remplace :</font> {esc(s.get('original', ''))}", old),
            Paragraph(f"<font color='#999999'>Par :</font> {esc(s.get('suggestion', ''))}", new),
        ]

    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                            topMargin=2.2 * cm, bottomMargin=2.2 * cm)
    doc.build(story)

    rel = str(out_path.relative_to(ROOT))
    db.update_job(job_id, tailored_cv_path=rel, status="cv_tailored")
    return out_path


def main(argv: list[str]) -> int:
    # Internal: background DOCX→PDF upgrade spawned by swap_docx. Silent.
    if argv and argv[0] == "__convert__":
        job_id, docx_rel = int(argv[1]), argv[2]
        pdf = docx_to_pdf(ROOT / docx_rel)
        if pdf is not None:
            db.update_job(job_id, tailored_cv_path=str(pdf.relative_to(ROOT)))
        return 0

    if len(argv) < 2:
        print(__doc__)
        return 1
    job_id = int(argv[0])
    if argv[1] == "--suggestions":
        path = write_suggestions(job_id, json.loads(argv[2]))
        print(f"Suggestions CV : {path}")
    else:
        path = swap_docx(job_id, json.loads(argv[1]))
        print(f"CV prêt (design préservé) — PDF en cours en arrière-plan : {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
